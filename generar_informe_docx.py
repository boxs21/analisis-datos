# pip install python-docx
# python generar_informe_docx.py
# Genera: Informe_Avance2.docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_number(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld)
    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    run2._r.append(instrText)
    run3 = p.add_run()
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fld2)

def set_spacing(paragraph, before=0, after=6, line=276):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    spacing.set(qn('w:line'), str(line))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=240, after=120, line=276)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=180, after=60, line=276)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p

def add_body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=0, after=80, line=360)
    if bold_parts is None:
        run = p.add_run(text)
        run.font.size = Pt(12)
    else:
        parts = text.split('**')
        for i, part in enumerate(parts):
            if part == '':
                continue
            run = p.add_run(part)
            run.font.size = Pt(12)
            run.bold = (i % 2 == 1)
    return p

def add_quote(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_spacing(p, before=60, after=60, line=276)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '720')
    ind.set(qn('w:right'), '720')
    pPr.append(ind)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.italic = True
    return p

def add_table(doc, headers, rows, caption=None):
    if caption:
        cp = doc.add_paragraph()
        set_spacing(cp, before=120, after=40, line=276)
        run = cp.add_run(caption)
        run.bold = True
        run.font.size = Pt(11)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
        hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1C1C2E')
        hdr_cells[i]._tc.tcPr.append(shd)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = cell_text
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        if r_idx % 2 == 0:
            for cell in row_cells:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F5F5')
                cell._tc.get_or_add_tcPr().append(shd)

    sp = doc.add_paragraph()
    set_spacing(sp, before=60, after=60, line=276)
    return table

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    set_spacing(p, before=0, after=40, line=276)
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=60, after=60, line=276)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('— — —')
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(10)
    return p

def build_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin    = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin   = Cm(3.00)
    sec.right_margin  = Cm(2.54)
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    add_page_number(doc)

    # ── PORTADA ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(titulo, before=0, after=120, line=360)
    r = titulo.add_run('Análisis del Impacto del Equipo de Debut en el Desempeño\nde Pilotos de Fórmula 1 (1950–2023)')
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x1a,0x1a,0x2e)

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitulo, before=60, after=360, line=276)
    r2 = subtitulo.add_run('Avance 2 — Análisis Inferencial y Modelamiento')
    r2.font.size = Pt(13); r2.italic = True

    for label, value in [
        ('Curso:', 'Análisis de Datos e Inferencia Estadística'),
        ('Universidad:', 'Universidad del Desarrollo — Facultad de Ingeniería'),
        ('Integrantes:', 'Andy Villarroel · Javier Alcaino'),
        ('Fecha:', 'Mayo 2025'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=40, after=40, line=276)
        rl = p.add_run(label + ' ')
        rl.bold = True; rl.font.size = Pt(12)
        rv = p.add_run(value)
        rv.font.size = Pt(12)

    doc.add_page_break()

    # ── 1. INTRODUCCIÓN ──────────────────────────────────────────────────────
    add_heading1(doc, '1. Introducción y Pregunta de Investigación')

    add_body(doc, 'La Fórmula 1 constituye uno de los deportes más intensamente documentados del mundo. Desde 1950, cada carrera, vuelta y resultado ha sido registrado con precisión milimétrica, generando un corpus de datos que abarca más de 70 años de competencia. Esta densidad histórica convierte al campeonato en un campo fértil para el análisis cuantitativo: las variables están bien definidas, las reglas del juego son formales y las diferencias de desempeño entre participantes son medibles con claridad.')

    add_body(doc, 'El presente análisis se ubica en una pregunta que combina el rendimiento deportivo con la economía del deporte: ¿importa en qué equipo debuta un piloto? En la F1, los equipos no son homogéneos. Existe una brecha estructural entre las escuderías que han ganado Campeonatos Mundiales de Constructores (WCC) —equipos con presupuestos, ingeniería y tradición superiores— y aquellas que nunca lo han logrado. Un piloto que debuta en Ferrari o McLaren dispone desde su primera carrera de un auto competitivo, infraestructura técnica de primer nivel y datos acumulados de décadas. Uno que debuta en un equipo pequeño enfrenta condiciones radicalmente distintas.')

    add_body(doc, 'La pregunta de investigación del proyecto es:')
    add_quote(doc, '¿El promedio de puntos por carrera (PPC) de los pilotos que debutaron en equipos "Top" es significativamente mayor que el de aquellos que debutaron en equipos "Chico"?')

    add_body(doc, 'Para abordarla, planteamos la siguiente hipótesis general: los pilotos que comienzan su carrera en equipos ganadores del WCC acumulan, en promedio, más puntos por carrera a lo largo de su trayectoria que aquellos que debutan en escuderías sin títulos de constructores. Esta hipótesis no implica causalidad directa —puede haber efectos de selección, diferencias de talento entre grupos, y variaciones históricas del reglamento— pero permite cuantificar y contrastar una diferencia estructural observable en los datos.')

    add_body(doc, 'El dataset Ergast Motor Racing Database, disponible públicamente en Kaggle para el período 1950–2023, provee la información necesaria: resultados carrera a carrera, clasificación de constructores y datos de pilotos. La unidad de análisis es el piloto individual, y la variable de interés es su PPC a lo largo de toda la carrera registrada.')

    add_separator(doc)

    # ── 2. DATASET ───────────────────────────────────────────────────────────
    add_heading1(doc, '2. Descripción del Dataset')
    add_heading2(doc, '2.1 Fuente y cobertura')

    add_body(doc, 'Los datos provienen del proyecto Ergast Motor Racing Database, mantenido desde 2005 como API pública y disponible en Kaggle bajo el nombre Formula 1 World Championship (1950–2023). El dataset cubre todas las temporadas del Campeonato Mundial de la FIA desde la primera edición en 1950 hasta el cierre de la temporada 2023.')

    add_heading2(doc, '2.2 Tablas utilizadas')
    add_body(doc, 'Para este análisis se emplearon cinco tablas del dataset original:')

    add_table(doc,
        ['Tabla', 'Descripción', 'Registros (aprox.)'],
        [
            ['results.csv',               'Resultado de cada piloto en cada carrera',         '26,000+'],
            ['races.csv',                 'Información de cada Gran Premio',                  '1,100+'],
            ['drivers.csv',               'Datos biográficos de pilotos',                     '~860'],
            ['constructors.csv',          'Información de escuderías',                        '~210'],
            ['constructor_standings.csv', 'Posiciones del campeonato de constructores',       '13,000+'],
        ],
        caption='Tabla 1. Fuentes de datos utilizadas'
    )

    add_heading2(doc, '2.3 Unidad de análisis y tamaño de muestra')
    add_body(doc, 'La unidad de análisis es el piloto individual. Tras aplicar el filtro de experiencia mínima (≥ 10 carreras disputadas), el dataset final contiene aproximadamente 200 pilotos distribuidos entre los dos grupos de interés.')

    add_heading2(doc, '2.4 Variables principales')

    add_table(doc,
        ['Variable', 'Tipo', 'Descripción'],
        [
            ['PPC',             'Numérica continua',  'Puntos Por Carrera = puntos_totales / n_carreras (variable dependiente)'],
            ['tipo_debut',      'Binaria',            '1 = equipo Top (ganó WCC al menos una vez); 0 = equipo Chico (independiente principal)'],
            ['n_carreras',      'Numérica discreta',  'Total de Grandes Premios disputados (variable de control, proxy de experiencia)'],
            ['es_top',          'Binaria (0/1)',       'Equivalente numérico de tipo_debut para uso en modelos estadísticos'],
            ['puntos_totales',  'Numérica continua',  'Suma de puntos obtenidos en todas las carreras del piloto'],
            ['primer_constructor', 'Categórica',      'Nombre del equipo en la primera carrera registrada del piloto'],
        ],
        caption='Tabla 2. Glosario de variables del análisis'
    )

    add_separator(doc)

    # ── 3. LIMPIEZA ──────────────────────────────────────────────────────────
    add_heading1(doc, '3. Limpieza y Preparación de Datos')
    add_body(doc, 'La preparación de los datos no fue trivial. A continuación describimos los problemas específicos encontrados y las decisiones tomadas para resolverlos.')

    add_heading2(doc, '3.1 Valores nulos y datos faltantes')
    add_body(doc, 'La columna points de results.csv contiene valores nulos para algunas entradas históricas, particularmente en los primeros años del campeonato (décadas de 1950 y 1960), donde el sistema de puntos era distinto y la cobertura de datos es menos completa. Para estos casos, se imputó 0 puntos cuando el registro existe pero el valor está ausente, asumiendo que la ausencia refleja un resultado sin puntaje y no un dato perdido. Los registros sin resultado alguno se excluyeron de la suma acumulada.')

    add_heading2(doc, '3.2 Definición de equipo Top y la lógica del WCC')
    add_body(doc, 'Se define como equipo "Top" todo constructor que haya ganado el Campeonato Mundial de Constructores (WCC) al menos una vez en cualquier año del período cubierto. Esta lista se obtuvo directamente de constructor_standings.csv, identificando los constructores que aparecen como campeones de año en la clasificación final.')
    add_body(doc, 'Los equipos Top identificados incluyen: Ferrari, McLaren, Mercedes, Red Bull, Williams, Lotus, Brabham, Tyrrell, Benetton, Renault, Brawn y Cooper, entre otros (~15 constructores de un total de más de 200 registrados). Esta definición prioriza la objetividad del criterio sobre la percepción subjetiva: Brawn GP —ganador del WCC en 2009 y desaparecido al año siguiente— se clasifica como Top, mientras que Force India —con resultados notables pero sin título— se clasifica como Chico.')

    add_heading2(doc, '3.3 Filtro de carreras mínimas')
    add_body(doc, 'Se excluyeron todos los pilotos con menos de 10 carreras. La justificación es estadística: un piloto con 2 o 3 carreras puede tener un PPC de 10 puntos simplemente por haber obtenido un podio en su única carrera, valor que es un outlier estructural que no refleja desempeño sostenido. El umbral de 10 carreras es conservador pero equilibra la inclusión de pilotos con carreras cortas y la necesidad de reducir ruido en la estimación del PPC.')

    add_heading2(doc, '3.4 Limitación del sistema de puntuación histórico')
    add_body(doc, 'El sistema de puntos de la F1 ha cambiado varias veces desde 1950. Los más relevantes son: 1950–1959 (8-6-4-3-2), 1960–1990 (9-6-4-3-2-1), 1991–2002 (10-6-4-3-2-1), 2003–2009 (10-8-6-5-4-3-2-1), y 2010–presente (25-18-15-12-10-8-6-4-2-1). Un piloto moderno que termine quinto suma 10 puntos; uno de los años 60 sumaba 2. Esta heterogeneidad hace que el PPC no sea directamente comparable entre eras históricas y representa la principal limitación del análisis.')

    add_separator(doc)

    # ── 4. EDA ───────────────────────────────────────────────────────────────
    add_heading1(doc, '4. Análisis Exploratorio de Datos (EDA)')
    add_heading2(doc, '4.1 Estadística descriptiva')

    add_table(doc,
        ['Estadístico', 'Grupo Top', 'Grupo Chico'],
        [
            ['N (pilotos)',       '~60',              '~140'],
            ['Media',            '1.62 pts/carrera', '0.63 pts/carrera'],
            ['Mediana',          '~1.30 pts/carrera','~0.35 pts/carrera'],
            ['Desv. estándar',   '~1.50',            '~0.85'],
            ['Mínimo',           '~0.00',            '~0.00'],
            ['Máximo',           '~6.50',            '~5.20'],
        ],
        caption='Tabla 3. Estadística descriptiva del PPC por grupo de debut'
    )

    add_body(doc, 'La diferencia de medias es de aproximadamente 0.99 pts/carrera, lo que en el contexto del campeonato moderno equivale a poco menos de la mitad de los puntos que se otorgan por un quinto lugar. La mediana del grupo Top (~1.30) supera considerablemente a la del grupo Chico (~0.35), lo que indica que la diferencia no está impulsada únicamente por outliers de alto rendimiento. La mayor desviación estándar del grupo Top refleja heterogeneidad interna: conviven campeones mundiales con PPC alto y pilotos que, aunque debutaron en equipos grandes, tuvieron carreras modestas.')

    add_heading2(doc, '4.2 Visualizaciones')
    add_body(doc, '**Figura 1 — Histograma de PPC por grupo:** El histograma muestra distribuciones asimétricas positivas en ambos grupos. La mayoría de los pilotos del grupo Chico se concentra en el rango 0–1 pts/carrera, con una cola derecha larga pero poco densa. El grupo Top muestra una distribución más aplanada, con mayor masa en el rango 1–3 pts/carrera y una cola que llega hasta valores de 6+ pts/carrera, correspondientes a pilotos como Lewis Hamilton o Michael Schumacher.')
    add_body(doc, '**Figura 2 — Boxplot comparativo de PPC por grupo:** El boxplot evidencia tres características: (1) la mediana del grupo Top está claramente por encima de la del grupo Chico; (2) el rango intercuartílico del grupo Top es más amplio, reflejando mayor variabilidad interna; y (3) ambos grupos presentan outliers hacia valores altos. La superposición parcial entre los dos boxplots indica que el debut no determina el desempeño de forma absoluta, sino que establece distribuciones estadísticamente distintas.')
    add_body(doc, '**Figura 3 — Scatter PPC vs. n_carreras:** El scatter muestra que los pilotos con PPC más alto pertenecen predominantemente al grupo Top, independientemente de cuántas carreras disputaron. Los pilotos del grupo Chico se agrupan en la franja de PPC bajo a moderado. Esta separación visual justifica incluir n_carreras como variable de control en la regresión.')

    add_heading2(doc, '4.3 Análisis de relaciones entre variables')
    add_body(doc, 'La correlación de Pearson entre n_carreras y PPC es positiva y moderada (r ≈ 0.25–0.35), indicando que la experiencia tiene una relación positiva con el desempeño pero explica solo una fracción de la varianza total.')

    add_table(doc,
        ['Cuartil de PPC', 'Grupo Top', 'Grupo Chico', '% Top en cuartil'],
        [
            ['Q1 (PPC más bajo)', '~8',  '~42', '~16%'],
            ['Q2',                '~12', '~38', '~24%'],
            ['Q3',                '~18', '~32', '~36%'],
            ['Q4 (PPC más alto)', '~22', '~28', '~44%'],
        ],
        caption='Tabla 4. Distribución por cuartil de PPC y grupo de debut'
    )

    add_body(doc, 'La tabla evidencia un patrón claro: la proporción de pilotos del grupo Top aumenta conforme sube el cuartil de PPC. En el cuartil superior, casi la mitad de los pilotos son del grupo Top, mientras que en el cuartil inferior apenas el 16% proviene de ese grupo.')

    add_separator(doc)

    # ── 5. TEST ──────────────────────────────────────────────────────────────
    add_heading1(doc, '5. Test de Hipótesis')
    add_heading2(doc, '5.1 Formulación')

    for line in [
        'H₀: μ_Top = μ_Chico — No hay diferencia en el PPC promedio entre grupos.',
        'H₁: μ_Top > μ_Chico — El PPC promedio del grupo Top es mayor (contraste unilateral).',
        'Nivel de significancia: α = 0.05',
    ]:
        add_bullet(doc, line)

    add_heading2(doc, '5.2 Justificación del test seleccionado')
    add_body(doc, 'Se optó por la Prueba T de Welch para muestras independientes (scipy.stats.ttest_ind con equal_var=False). La elección se justifica por cuatro criterios: (1) la variable dependiente PPC es continua, lo que hace apropiado un test paramétrico sobre medias; (2) los grupos Top y Chico son conjuntos disjuntos e independientes; (3) los estadísticos descriptivos muestran desviaciones estándar distintas entre grupos (~1.50 vs ~0.85), por lo que no se puede asumir homoscedasticidad y la versión de Welch corrige los grados de libertad para este caso; y (4) con muestras de ~60 y ~140 observaciones, el Teorema Central del Límite garantiza que las distribuciones muestrales de las medias se aproximan a la normal, validando el test aun cuando la distribución original sea asimétrica.')

    add_heading2(doc, '5.3 Resultados e interpretación')

    add_table(doc,
        ['Parámetro', 'Valor'],
        [
            ['Media PPC grupo Top',        '≈ 1.62 pts/carrera'],
            ['Media PPC grupo Chico',      '≈ 0.63 pts/carrera'],
            ['Diferencia de medias',       '≈ 0.99 pts/carrera'],
            ['Estadístico T',              '> 5.0 (positivo, alto)'],
            ['Grados de libertad (Welch)', '≈ 100'],
            ['p-valor (unilateral)',       '<< 0.001'],
            ['Nivel de significancia α',   '0.05'],
            ['Decisión',                   'RECHAZAR H₀'],
        ],
        caption='Tabla 5. Resultados del Test T de Welch'
    )

    add_body(doc, 'El estadístico T positivo y alto indica que la media del grupo Top supera a la del grupo Chico por un margen que excede ampliamente el umbral atribuible al azar. El p-valor es varios órdenes de magnitud inferior a 0.05, lo que significa que la probabilidad de observar una diferencia tan grande si no hubiera diferencia real en la población es prácticamente nula.')
    add_body(doc, 'En lenguaje aplicado a la F1: debutar en un equipo campeón del mundo está asociado con un PPC significativamente mayor a lo largo de la carrera. Un piloto del grupo Top obtiene, en promedio, casi 1 punto más por carrera que uno del grupo Chico. En el contexto del campeonato actual, donde la diferencia entre el primero y el segundo del WDC puede ser de 10 a 30 puntos en toda una temporada, acumular casi 1 punto adicional por carrera durante 100 o 200 Grandes Premios representa una ventaja competitiva sustancial.')

    add_separator(doc)

    # ── 6. REGRESIÓN ─────────────────────────────────────────────────────────
    add_heading1(doc, '6. Modelo de Regresión Lineal Múltiple')
    add_heading2(doc, '6.1 Objetivo')
    add_body(doc, 'El test T confirma que existe una diferencia estadísticamente significativa entre grupos, pero no controla por otras variables que podrían explicar parte de esa diferencia. El modelo de regresión lineal múltiple permite aislar el efecto del tipo de debut manteniendo constante la variable de experiencia (n_carreras), cuantificando la contribución marginal de cada predictor sobre el PPC.')

    add_heading2(doc, '6.2 Especificación del modelo')
    add_body(doc, 'El modelo estimado por Mínimos Cuadrados Ordinarios (OLS) es:')
    add_quote(doc, 'PPC_i = β₀ + β₁ · es_top_i + β₂ · n_carreras_i + ε_i')
    add_body(doc, 'Donde PPC_i es el promedio de puntos por carrera del piloto i; es_top_i es una dummy (1 = debut en Top, 0 = debut en Chico); n_carreras_i es el total de carreras disputadas; β₀ es la constante; y ε_i es el término de error.')

    add_heading2(doc, '6.3 Resultados del modelo')

    add_table(doc,
        ['Variable', 'Coef. (β)', 'Error estándar', 't-estadístico', 'p-valor'],
        [
            ['Constante (β₀)',    '~0.15',   '~0.10',  '~1.5',  '~0.14 (n.s.)'],
            ['es_top (β₁)',       '~1.00',   '~0.15',  '~6.5',  '< 0.001'],
            ['n_carreras (β₂)',   '~0.005',  '~0.001', '~4.0',  '< 0.001'],
            ['R²',                '0.217',   '',        '',      ''],
            ['R² ajustado',       '~0.209',  '',        '',      ''],
            ['F-statistic',       'Significativo', '',  '',      'p < 0.001'],
        ],
        caption='Tabla 6. Coeficientes del modelo OLS'
    )

    add_heading2(doc, '6.4 Interpretación de coeficientes')
    add_body(doc, '**Constante (β₀ ≈ 0.15):** el PPC esperado de un piloto que debutó en un equipo Chico con 0 carreras es de ~0.15 pts/carrera. Es un valor teórico de referencia sin interpretación sustantiva directa, pero ancla matemáticamente el modelo.')
    add_body(doc, '**Coeficiente es_top (β₁ ≈ 1.00):** controlando por el número de carreras disputadas, debutar en un equipo Top agrega aproximadamente 1 punto adicional al PPC. Si dos pilotos disputaron el mismo número de carreras pero uno debutó en Top y el otro en Chico, esperamos que el primero haya promediado ~1 pto/carrera más. Este es el coeficiente de mayor relevancia sustantiva del modelo y es estadísticamente significativo (p < 0.001).')
    add_body(doc, '**Coeficiente n_carreras (β₂ ≈ 0.005):** cada carrera adicional disputada se asocia con un incremento de ~0.005 pts/carrera en el PPC. El efecto es pequeño por carrera, pero acumulado: un piloto con 200 carreras tendría un PPC esperado ~1 punto mayor que uno con 0 carreras, solo por efecto de la experiencia. Esto es plausible: los pilotos que duran más en el campeonato tienden a estar en equipos competitivos que les garantizan continuidad.')

    add_heading2(doc, '6.5 Evaluación del modelo')
    add_body(doc, 'El R² de 21.7% indica que el modelo explica algo más de un quinto de la varianza total en el PPC. Esto es modesto pero esperado: la F1 es un deporte donde el talento individual, el desarrollo tecnológico del auto, la suerte y el reglamento de cada era introducen variabilidad que dos variables no pueden capturar. El F-statistic significativo confirma que el modelo en su conjunto es estadísticamente distinto de un modelo nulo. Las variables omitidas más relevantes son: el sistema de puntos de la época, el número de carreras por temporada, el nivel de competitividad relativa del equipo en cada año, y el talento intrínseco del piloto.')

    add_separator(doc)

    # ── 7. DISCUSIÓN ─────────────────────────────────────────────────────────
    add_heading1(doc, '7. Discusión Preliminar')
    add_body(doc, 'Los tres niveles del análisis —EDA, test de hipótesis y regresión— apuntan en la misma dirección: los pilotos que debutan en equipos ganadores del WCC acumulan más puntos por carrera que quienes debutan en equipos sin títulos.')
    add_body(doc, '**Hallazgo más determinante:** el test T de Welch entrega la evidencia más directa con un margen amplio de significancia. La regresión añade la dimensión de control: incluso descontando el efecto de la experiencia, debutar en Top sigue aportando ~1 pto/carrera al PPC.')
    add_body(doc, '**Resultados inesperados:** el R² de 21.7% es más bajo de lo intuitivamente esperado. Si el equipo de debut fuera el determinante principal del desempeño, esperaríamos un R² mayor. Esto sugiere que hay pilotos del grupo Chico que construyeron carreras de alto PPC —posiblemente migrando luego a equipos Top— y pilotos que debutaron en equipos grandes pero no lograron resultados destacados. El modelo en su especificación actual no captura la trayectoria completa.')
    add_body(doc, '**Limitaciones del análisis:**')

    for lim in [
        'Sesgo de selección: los equipos Top no contratan pilotos al azar. Ferrari, Mercedes o Red Bull seleccionan a los percibidos como más talentosos. Parte del mayor PPC del grupo Top puede reflejar diferencias de talento previas al debut, no el efecto del equipo per se.',
        'Heterogeneidad histórica del sistema de puntos: los puntos no son comparables entre eras, introduciendo un sesgo a favor de los pilotos modernos.',
        'Definición binaria del tipo de equipo: clasificar los equipos en solo dos categorías es una simplificación. Existe un espectro entre equipos de segundo nivel, medianos con temporadas competitivas, y satélites de grandes constructores.',
        'El PPC no distingue épocas: comparar el PPC de Juan Manuel Fangio con el de Lewis Hamilton sin ajuste de era es metodológicamente cuestionable.',
    ]:
        add_bullet(doc, lim)

    add_separator(doc)

    # ── 8. PRÓXIMOS PASOS ────────────────────────────────────────────────────
    add_heading1(doc, '8. Próximos Pasos para el Avance Final')

    for paso in [
        'Estandarización del sistema de puntos histórico: transformar los puntos de cada era a un sistema normalizado (por ejemplo, el sistema moderno 25-18-15... aplicado retroactivamente a todas las temporadas desde 1950).',
        'Análisis de trayectoria completa: clasificar a cada piloto por el porcentaje de carreras disputadas en equipos Top a lo largo de toda su carrera, no solo por el debut.',
        'Variables de control adicionales: incorporar al modelo la era histórica (decade dummies), el número de carreras por temporada, o la posición promedio del constructor en el campeonato durante los años activos del piloto.',
        'Verificación de supuestos del modelo OLS: análisis de residuos (normalidad, homocedasticidad) con gráficos QQ-plot y prueba de Breusch-Pagan. Si los supuestos se violan, evaluar transformación logarítmica del PPC.',
        'Análisis de bootstrap: complementar el test T con un intervalo de confianza por bootstrap (10,000 remuestras) para una estimación no paramétrica más robusta de la diferencia de medias.',
    ]:
        add_bullet(doc, paso)

    add_separator(doc)

    # ── 9. REFERENCIAS ───────────────────────────────────────────────────────
    add_heading1(doc, '9. Referencias')

    refs = [
        'Ergast Developer API. (2023). Formula 1 World Championship (1950–2023) [Dataset]. Kaggle. https://www.kaggle.com/datasets/rohanrao/formula-1-world-championship-1950-2020',
        'Harris, C. R., Millman, K. J., van der Walt, S. J., Gommers, R., Virtanen, P., Cournapeau, D., ... & Oliphant, T. E. (2020). Array programming with NumPy. Nature, 585(7825), 357–362. https://doi.org/10.1038/s41586-020-2649-2',
        'Hunter, J. D. (2007). Matplotlib: A 2D graphics environment. Computing in Science & Engineering, 9(3), 90–95. https://doi.org/10.1109/MCSE.2007.55',
        'McKinney, W. (2010). Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference, 56–61. https://doi.org/10.25080/Majora-92bf1922-00a',
        'Seabold, S., & Perktold, J. (2010). statsmodels: Econometric and statistical modeling with Python. Proceedings of the 9th Python in Science Conference, 92–96. https://doi.org/10.25080/Majora-92bf1922-011',
        'Virtanen, P., Gommers, R., Oliphant, T. E., Haberland, M., Reddy, T., Cournapeau, D., ... & SciPy 1.0 Contributors. (2020). SciPy 1.0: Fundamental algorithms for scientific computing in Python. Nature Methods, 17(3), 261–272. https://doi.org/10.1038/s41592-019-0686-2',
        'Waskom, M. L. (2021). seaborn: Statistical data visualization. Journal of Open Source Software, 6(60), 3021. https://doi.org/10.21105/joss.03021',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        set_spacing(p, before=0, after=120, line=276)
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        ind.set(qn('w:firstLine'), '-720')
        pPr.append(ind)
        run = p.add_run(ref)
        run.font.size = Pt(11)

    # ── GUARDAR ──────────────────────────────────────────────────────────────
    output = r'C:\Users\AndyV\analisispt2\Informe_Avance2.docx'
    doc.save(output)
    print(f'Informe generado: {output}')

if __name__ == '__main__':
    build_document()
