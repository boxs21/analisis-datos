# pip install python-docx
# python generar_guion_docx.py
# Genera: Guion_Presentacion.docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED   = RGBColor(0xE1, 0x06, 0x00)
DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1E, 0x88, 0x50)

def spacing(para, before=0, after=80, line=276):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(before)); sp.set(qn('w:after'), str(after))
    sp.set(qn('w:line'), str(line));     sp.set(qn('w:lineRule'), 'auto')
    pPr.append(sp)

def h1(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=320, after=100, line=276)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK
    return p

def h2(doc, text):
    p = doc.add_paragraph()
    spacing(p, before=200, after=60, line=276)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RED
    return p

def body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    spacing(p, before=0, after=80, line=320)
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        r.font.size = Pt(11)
        r.bold = (i % 2 == 1)
    return p

def speaker(doc, name, slide_title, tiempo):
    p = doc.add_paragraph()
    spacing(p, before=240, after=20, line=276)
    r1 = p.add_run(f'[{name}]  ')
    r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = RED if name == 'ANDY' else RGBColor(0x1A, 0x6B, 0xD0)
    r2 = p.add_run(slide_title)
    r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = DARK
    r3 = p.add_run(f'  ·  {tiempo}')
    r3.font.size = Pt(10); r3.font.color.rgb = GRAY
    return p

def nota(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    spacing(p, before=40, after=60, line=276)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), '480')
    pPr.append(ind)
    r = p.add_run(f'[Nota: {text}]')
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY
    return p

def pregunta(doc, num, pregunta_text, respuesta_text):
    p = doc.add_paragraph()
    spacing(p, before=240, after=60, line=276)
    r = p.add_run(f'Pregunta {num}: {pregunta_text}')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK
    body(doc, respuesta_text)

def separador(doc):
    p = doc.add_paragraph()
    spacing(p, before=80, after=80, line=276)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('— — —')
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC); r.font.size = Pt(10)

def tabla_tiempos(doc, rows):
    t = doc.add_table(rows=1 + len(rows), cols=3)
    t.style = 'Table Grid'
    hdrs = ['Slide', 'Presentador', 'Tiempo estimado']
    for i, h in enumerate(hdrs):
        cell = t.rows[0].cells[i]
        cell.text = h
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True; run.font.size = Pt(10)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A1A2E')
        cell._tc.get_or_add_tcPr().append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for idx, (s, p_name, t_val) in enumerate(rows):
        row = t.rows[idx + 1].cells
        for ci, val in enumerate([s, p_name, t_val]):
            row[ci].text = val
            for para in row[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    # ── PORTADA ────────────────────────────────────────────────────
    doc.add_paragraph(); doc.add_paragraph()

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(titulo, before=0, after=120, line=360)
    r = titulo.add_run('GUION DE PRESENTACIÓN ORAL')
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RED

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(subtitulo, before=60, after=60, line=276)
    r2 = subtitulo.add_run('¿Debutar en un equipo grande te da más puntos en la F1?')
    r2.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = DARK

    doc.add_paragraph()
    for label, value in [
        ('Curso:', 'Análisis de Datos e Inferencia Estadística'),
        ('Universidad:', 'Universidad del Desarrollo'),
        ('Integrantes:', 'Andy Villarroel  ·  Javier Alcaino'),
        ('Duración:', '~12 minutos  (14 slides)'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, before=30, after=30, line=276)
        rl = p.add_run(label + '  '); rl.bold = True; rl.font.size = Pt(11)
        rv = p.add_run(value); rv.font.size = Pt(11)

    doc.add_paragraph()
    aviso = doc.add_paragraph()
    aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(aviso, before=60, after=60, line=276)
    r3 = aviso.add_run('INSTRUCCIONES: [ANDY] y [JAVIER] indican quién habla. Textos entre corchetes son acotaciones escénicas — no se leen en voz alta.')
    r3.italic = True; r3.font.size = Pt(10); r3.font.color.rgb = GRAY

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # BLOQUE ANDY — Slides 1 a 7
    # ═══════════════════════════════════════════════════════════════

    # ── SLIDE 1 ────────────────────────────────────────────────────
    speaker(doc, 'ANDY', 'Slide 1 — Portada', '0:40')
    body(doc, 'Buenas tardes. Somos Andy Villarroel y Javier Alcaino, y hoy les presentamos nuestro análisis de datos sobre Fórmula 1. El título del proyecto es: *¿Debutar en un equipo grande te da más puntos en la F1?*')
    body(doc, 'En los próximos doce minutos les vamos a contar qué pregunta nos hicimos, cómo construimos el análisis, qué encontramos y por qué creemos que los resultados son interesantes aunque, como verán, tienen sus matices.')
    nota(doc, 'Andy señala el título en la pantalla. Javier está de pie al lado, asiente. Pausa breve antes de pasar.')

    separador(doc)

    # ── SLIDE 2 ────────────────────────────────────────────────────
    speaker(doc, 'ANDY', 'Slide 2 — Agenda', '0:30')
    body(doc, 'La presentación está dividida en cuatro bloques. Primero les explicamos la pregunta de investigación y el dataset. Luego mostramos el análisis exploratorio de datos. Después Javier les presenta las pruebas estadísticas formales: el test de hipótesis y la regresión. Y cerramos con la discusión, limitaciones y próximos pasos.')
    nota(doc, 'Andy recorre visualmente los ítems de la agenda. No leer cada ítem palabra por palabra, hablar fluidamente.')

    separador(doc)

    # ── SLIDE 3 ────────────────────────────────────────────────────
    speaker(doc, 'ANDY', 'Slide 3 — Pregunta de Investigación', '1:00')
    body(doc, 'Todo buen análisis parte de una pregunta concreta. La nuestra es esta: **¿el promedio de puntos por carrera de pilotos que debutaron en equipos Top es significativamente mayor que el de pilotos que debutaron en equipos chicos?**')
    body(doc, 'Para ser precisos con los términos: definimos "equipo Top" como cualquier equipo que haya ganado al menos un Campeonato Mundial de Constructores —el WCC. Eso incluye a Ferrari, McLaren, Mercedes, Red Bull, Williams, Renault, Lotus, entre otros. Equipos que han levantado el trofeo al menos una vez.')
    body(doc, 'Y nuestra métrica de rendimiento es el **Promedio de Puntos por Carrera**, que abreviamos PPC. Lo calculamos dividiendo el total de puntos acumulados por un piloto en toda su carrera entre el número de carreras que disputó.')
    body(doc, 'La hipótesis nula es que no hay diferencia entre grupos. La hipótesis alternativa es que los pilotos de equipos Top tienen un PPC mayor. En un momento Javier les explica por qué tomamos esa dirección y qué test usamos.')
    nota(doc, 'Señalar la fórmula del PPC en la slide. Tomarse este slide con calma — el marco conceptual es clave y el profe prestará mucha atención aquí.')

    separador(doc)

    # ── SLIDE 4 ────────────────────────────────────────────────────
    speaker(doc, 'ANDY', 'Slide 4 — Dataset', '1:10')
    body(doc, 'Para responder esa pregunta usamos el dataset de Ergast Motor Racing, que es una API pública con datos históricos de Fórmula 1 desde 1950 hasta la temporada 2023. Es el dataset de referencia para este tipo de análisis.')
    body(doc, 'El dataset incluye múltiples tablas: resultados por carrera, información de pilotos, datos de constructores y clasificaciones de campeonato. Nosotros trabajamos principalmente con las tablas de resultados, constructores y clasificaciones de constructores.')
    body(doc, 'En total, después de la limpieza que les muestro enseguida, trabajamos con **377 pilotos**: **123 del grupo Top** y **254 del grupo Chico**. El dataset cubre más de setenta años de historia del deporte, lo que da una variabilidad enorme — y esa variabilidad es parte de lo que hace el análisis interesante.')
    nota(doc, 'Señalar la tabla del dataset y los dos comparadores de grupo en la slide.')

    separador(doc)

    # ── SLIDE 5 ────────────────────────────────────────────────────
    speaker(doc, 'ANDY', 'Slide 5 — Limpieza de Datos', '1:00')
    body(doc, 'Los datos crudos no llegaron listos para usar. Tuvimos que tomar varias decisiones metodológicas.')
    body(doc, 'El criterio más importante fue **filtrar pilotos con menos de 10 carreras disputadas**. La razón es estadística: si un piloto corrió solo dos o tres carreras, su PPC no es representativo de su rendimiento real. Podría ser alguien que ganó una carrera en su debut y no volvió. Con menos de 10 carreras hay demasiado ruido.')
    body(doc, 'También identificamos el equipo de debut de cada piloto — el constructor con el que corrió su primera carrera — y lo clasificamos como Top o Chico según la definición ya mencionada.')
    body(doc, 'Además manejamos valores nulos en puntos, que corresponden a carreras donde el piloto se retiró o no clasificó. Esos se tratan como cero puntos, lo cual es lo más honesto metodológicamente.')
    nota(doc, 'Mantener el ritmo, este slide no requiere demasiado detalle técnico salvo que el profe pregunte.')
    body(doc, '**[Transición]:** Con los datos limpios, pasamos a explorarlos.')

    separador(doc)

    # ── SLIDE 6 ────────────────────────────────────────────────────
    speaker(doc, 'ANDY', 'Slide 6 — EDA: Estadística Descriptiva', '1:00')
    body(doc, 'En el análisis exploratorio, lo primero que hacemos es describir los datos. Y los números ya cuentan una historia bastante clara.')
    body(doc, 'El grupo de pilotos que debutaron en equipos Top tiene un **PPC promedio de 1.62 puntos por carrera**. El grupo que debutó en equipos chicos tiene un promedio de **0.63 puntos por carrera**. Eso es una diferencia de casi 1 punto por carrera, lo que en Fórmula 1 puede significar la diferencia entre terminar sexto o décimo en el campeonato al final de la temporada.')
    body(doc, 'Las medianas cuentan una historia similar. Y las desviaciones estándar nos dicen que ambos grupos tienen bastante dispersión interna — dentro de los equipos Top también hay pilotos que no llegaron a nada, y en los equipos chicos también hay pilotos que se destacaron enormemente.')
    nota(doc, 'Señalar los valores de la tabla descriptiva en la pantalla mientras los nombra.')

    separador(doc)

    # ── SLIDE 7 ────────────────────────────────────────────────────
    speaker(doc, 'ANDY', 'Slide 7 — Histograma de Distribución', '0:50')
    body(doc, 'El histograma nos muestra la distribución completa del PPC para ambos grupos superpuestos.')
    body(doc, 'Lo que vemos es que ambas distribuciones tienen asimetría positiva: la mayoría de los pilotos se acumula en valores bajos de PPC y hay una cola larga hacia la derecha. Eso es esperable — los pilotos de alto rendimiento son excepcionales por definición.')
    body(doc, 'La distribución del grupo Top está claramente corrida hacia la derecha respecto al grupo Chico. La concentración del grupo Chico está muy cercana a cero, mientras que el grupo Top tiene mayor masa en rangos intermedios y altos.')
    body(doc, 'Esta distribución sesgada es, dicho sea de paso, una de las razones por las que vamos a usar el t-test de Welch para el test formal — pero eso se los explica Javier.')
    nota(doc, 'Señalar en la pantalla la diferencia de centros entre ambas distribuciones. Este es el cierre del bloque de Andy.')
    body(doc, '**[Transición a Javier]:** Con ese panorama descriptivo sobre la mesa, le paso la palabra a Javier para las pruebas estadísticas formales.')

    separador(doc)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # BLOQUE JAVIER — Slides 8 a 14
    # ═══════════════════════════════════════════════════════════════

    # ── SLIDE 8 ────────────────────────────────────────────────────
    speaker(doc, 'JAVIER', 'Slide 8 — Boxplot y Scatter', '1:10')
    body(doc, 'Gracias, Andy. Antes de entrar al test de hipótesis, quiero mostrarles dos visualizaciones más.')
    body(doc, 'El **boxplot** confirma visualmente lo que Andy describió: la mediana del grupo Top es notablemente más alta que la del grupo Chico, y la caja del grupo Top está completamente por encima. También vemos más outliers en el grupo Top — esos corresponden a las leyendas del deporte: Hamilton, Schumacher, Senna.')
    body(doc, 'El **scatter plot** muestra el PPC versus el número de carreras disputadas. Aquí hay algo interesante: hay una leve tendencia positiva. Los pilotos con más carreras tienden a tener PPC más altos. Los pilotos que duran más en la grilla son los que más puntúan. Eso ya nos da el indicio de que el número de carreras podría ser una covariable relevante, algo que vamos a incorporar en la regresión.')
    nota(doc, 'Señalar primero el boxplot, luego el scatter. Si están en slides separadas, pasar rápido de una a otra.')

    separador(doc)

    # ── SLIDE 9 ────────────────────────────────────────────────────
    speaker(doc, 'JAVIER', 'Slide 9 — Test de Hipótesis', '1:30')
    body(doc, 'Ahora sí, el test formal. Nuestras hipótesis son:')
    body(doc, '**H₀:** el PPC promedio de ambos grupos es igual — no hay diferencia.  **H₁:** el PPC promedio del grupo Top es mayor que el del grupo Chico.')
    body(doc, 'Usamos el **test T de Welch** —también llamado T con varianzas desiguales. ¿Por qué Welch y no el T estándar? Porque el test de Student asume que las varianzas de ambos grupos son iguales. Nosotros no tenemos razón para asumir eso: los dos grupos tienen tamaños distintos —123 y 254— y dispersiones distintas, como vimos en los boxplots. El test de Welch es más robusto en esa situación y es el estándar recomendado. Por eso en la implementación usamos equal_var=False.')
    body(doc, 'El resultado fue contundente: el **estadístico T fue 4.61** y el **p-valor de 7.86 × 10⁻⁶**, varios órdenes de magnitud por debajo de 0.05. Rechazamos la hipótesis nula con un nivel de confianza muy alto. Hay evidencia estadística sólida de que los pilotos que debutaron en equipos Top tienen un PPC significativamente mayor.')
    body(doc, 'Ahora bien, rechazar H₀ no nos dice cuánto mayor es la diferencia ni si hay otras variables involucradas. Para eso necesitamos el modelo de regresión.')
    nota(doc, 'Señalar el p-valor en la slide explícitamente. Si el profe pregunta por la distribución bilateral vs unilateral: el p-valor que mostramos es bilateral, para test unilateral se divide por dos, pero la conclusión no cambia.')

    separador(doc)

    # ── SLIDE 10 ────────────────────────────────────────────────────
    speaker(doc, 'JAVIER', 'Slide 10 — Modelo de Regresión OLS', '1:10')
    body(doc, 'Para profundizar en la relación, construimos un modelo de regresión lineal por mínimos cuadrados ordinarios — OLS. La variable dependiente es el PPC de cada piloto. Las variables independientes son dos: **es_top**, que es una variable binaria que vale 1 si el piloto debutó en un equipo Top y 0 si no; y **n_carreras**, que es el total de carreras disputadas.')
    body(doc, 'Incluimos n_carreras porque el scatter sugería una relación positiva y porque queremos aislar el efecto del debut: no queremos que la diferencia entre grupos se explique solo por el hecho de que los pilotos de equipos Top tienden a tener carreras más largas.')
    body(doc, 'El modelo resultante tiene un **R² de 41.6%**. Ambos coeficientes son estadísticamente significativos con p < 0.001. La constante también es significativa con p ≈ 0.04, aunque su valor es negativo y no tiene interpretación práctica directa. Vamos al slide siguiente para la interpretación.')
    nota(doc, 'Señalar la tabla de resultados del modelo en la pantalla. Mencionar los valores p de los coeficientes.')

    separador(doc)

    # ── SLIDE 11 ────────────────────────────────────────────────────
    speaker(doc, 'JAVIER', 'Slide 11 — Interpretación de Coeficientes', '1:10')
    body(doc, 'Interpretemos los coeficientes uno por uno.')
    body(doc, 'La **constante β₀ ≈ -0.21**: es el PPC teórico para un piloto de equipo Chico con cero carreras. Al ser negativo y fuera del rango real de los datos, no tiene interpretación sustantiva directa — simplemente ancla el modelo matemáticamente.')
    body(doc, 'El coeficiente **es_top ≈ +0.52** es el más importante: **manteniendo constante el número de carreras disputadas**, debutar en un equipo Top agrega en promedio **0.52 puntos adicionales al PPC**. Si dos pilotos disputaron exactamente las mismas carreras pero uno debutó en Top, esperamos que promedíe 0.52 pts/carrera más. Estadísticamente significativo con p < 0.001.')
    body(doc, 'Es importante notar que este 0.52 es distinto a la diferencia bruta de medias de 0.99 que vimos en el EDA. La regresión descuenta el efecto de la experiencia, y el efecto neto del debut es 0.52.')
    body(doc, 'El coeficiente **n_carreras ≈ +0.015**: cada carrera adicional suma 0.015 pts al PPC. Pequeño por carrera, pero acumulado: 200 carreras representan aproximadamente 3 puntos extra de PPC.')
    body(doc, 'Ahora bien, el **R² de 41.6%** indica que nuestras dos variables explican el 41.6% de la varianza total en el PPC. El resto lo explican factores fuera del modelo: el talento individual, el año en que corrió, el auto específico, la suerte en carrera. Eso no invalida el modelo — simplemente indica que el fenómeno es complejo y hemos capturado una parte relevante.')
    nota(doc, 'Señalar los coeficientes en la tabla mientras los interpreta. Distinguir explícitamente el 0.52 del 0.99 del EDA.')

    separador(doc)

    # ── SLIDE 12 ────────────────────────────────────────────────────
    speaker(doc, 'JAVIER', 'Slide 12 — Discusión y Limitaciones', '1:20')
    body(doc, 'Ahora lo más importante: ¿qué significa todo esto y cuáles son las limitaciones?')
    body(doc, 'Los resultados son claros en términos estadísticos. Hay una diferencia real y significativa en el PPC entre grupos, y persiste al controlar por experiencia. Pero la interpretación causal es delicada.')
    body(doc, 'La **limitación más grande** es el cambio histórico en el sistema de puntuación. En los años 50 el ganador se llevaba 9 puntos; hoy se llevan 25. Si un piloto corrió en los 60 y otro corrió en 2020, sus PPC simplemente no son comparables en escala. No estandarizamos los puntos en este avance — es la principal extensión que dejamos para el avance final.')
    body(doc, 'Una segunda limitación es la **confusión de variables**: ¿debutar en un equipo Top hace que el piloto sea mejor, o los equipos Top contratan a los pilotos que ya muestran ser mejores desde las categorías menores? No lo podemos determinar con este diseño observacional. Es una correlación, no una relación causal establecida.')
    body(doc, 'Dicho esto, el análisis cumple su objetivo: describe una diferencia real en los datos y la cuantifica con rigor estadístico.')
    nota(doc, 'Hablar con seguridad sobre las limitaciones. Mencionarlas no debilita el trabajo — al contrario, muestra pensamiento crítico. El profe valora esto.')

    separador(doc)

    # ── SLIDE 13 ────────────────────────────────────────────────────
    speaker(doc, 'JAVIER', 'Slide 13 — Próximos Pasos', '0:40')
    body(doc, 'Con estas limitaciones en mente, los próximos pasos naturales serían tres.')
    body(doc, 'Primero, **estandarizar los puntos según la época** — usar el porcentaje del máximo posible de puntos por temporada como métrica normalizada. Segundo, **incorporar la trayectoria completa** del piloto, no solo el debut: clasificar por el porcentaje de carreras en Top a lo largo de toda la carrera. Y tercero, **verificar los supuestos del OLS** con QQ-plot de residuos y test de Breusch-Pagan, y si fuera necesario, aplicar transformación logarítmica o bootstrap.')
    nota(doc, 'Este slide puede ir rápido. Es para mostrar que pensamos en el futuro del análisis.')

    separador(doc)

    # ── SLIDE 14 ────────────────────────────────────────────────────
    speaker(doc, 'JAVIER', 'Slide 14 — Conclusiones y Cierre', '0:30')
    body(doc, 'Para cerrar: encontramos evidencia estadística sólida de que los pilotos que debutan en equipos que han ganado el Campeonato de Constructores acumulan en promedio significativamente más puntos por carrera. La diferencia bruta es de casi 1 pto/carrera; el efecto controlado de la regresión es de 0.52 pts/carrera. Ambos son estadísticamente robustos.')
    body(doc, '"Si Hamilton hubiera debutado en un Minardi, ¿dónde estaría hoy? La respuesta estadística sería: con un PPC bastante más bajo. Pero eso es especulación — y nosotros somos gente de datos."')
    body(doc, '**Muchas gracias. Quedamos disponibles para preguntas.**')
    nota(doc, 'Ambos presentadores al frente. Javier termina el discurso. Dejar un silencio breve y natural después del cierre antes de abrir preguntas.')

    separador(doc)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # TABLA DE TIEMPOS
    # ═══════════════════════════════════════════════════════════════
    h1(doc, 'Resumen de Tiempos')
    tabla_tiempos(doc, [
        ('1 — Portada',                'Andy',   '0:40'),
        ('2 — Agenda',                 'Andy',   '0:30'),
        ('3 — Pregunta de Investigación','Andy', '1:00'),
        ('4 — Dataset',                'Andy',   '1:10'),
        ('5 — Limpieza de Datos',      'Andy',   '1:00'),
        ('6 — EDA Descriptiva',        'Andy',   '1:00'),
        ('7 — Histograma',             'Andy',   '0:50'),
        ('[Cambio de presentador]',    '',       ''),
        ('8 — Boxplot y Scatter',      'Javier', '1:10'),
        ('9 — Test de Hipótesis',      'Javier', '1:30'),
        ('10 — Regresión OLS',         'Javier', '1:10'),
        ('11 — Interpretación',        'Javier', '1:10'),
        ('12 — Discusión y Limitaciones','Javier','1:20'),
        ('13 — Próximos Pasos',        'Javier', '0:40'),
        ('14 — Conclusiones',          'Javier', '0:30'),
        ('TOTAL',                      '',       '~12:50'),
    ])

    separador(doc)
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # PREGUNTAS ANTICIPADAS
    # ═══════════════════════════════════════════════════════════════
    h1(doc, 'Sección de Preguntas Anticipadas')
    body(doc, 'Esta sección NO se lee en la presentación. Es para preparar respuestas durante la sesión de preguntas.')

    separador(doc)

    pregunta(doc, 1,
        '"¿Por qué no estandarizaron los puntos según la época?"',
        'Es una observación muy válida y es la principal limitación que reconocemos. No estandarizamos los puntos porque hacerlo requería definir un denominador de normalización — puede ser el máximo posible por temporada, la posición relativa en el campeonato, o un índice percentil. Todas esas opciones introducen supuestos propios. Optamos por reportar la limitación explícitamente en lugar de introducir una estandarización que podría no ser la más correcta sin un análisis adicional. Si tuviéramos que extender el trabajo, usaríamos el porcentaje de puntos obtenidos sobre el máximo posible en cada temporada como métrica normalizada.')

    separador(doc)

    pregunta(doc, 2,
        '"¿Cómo saben que el equipo causó el éxito y no el talento del piloto?"',
        'No lo sabemos, y somos explícitos en eso. Nuestro diseño es observacional, no experimental. No podemos afirmar causalidad en ninguna dirección. Lo que encontramos es una correlación estadísticamente significativa. La confusión entre talento y equipo es un problema clásico en análisis deportivos: los equipos buenos contratan pilotos buenos, y los pilotos buenos terminan en equipos buenos. Para separar esas variables necesitaríamos un diseño cuasi-experimental — por ejemplo, comparar pilotos que cambiaron de equipo Top a equipo chico o viceversa. Eso está fuera del alcance de este trabajo pero es el siguiente paso natural.')

    separador(doc)

    pregunta(doc, 3,
        '"El R² es 41.6% — ¿no indica que el modelo no sirve?"',
        'El R² de 41.6% con solo dos variables predictoras es un resultado sólido para un análisis de rendimiento deportivo. Estamos modelando el desempeño humano en un deporte que depende de cientos de factores: el auto, el motor, la estrategia, las condiciones climáticas, los accidentes, la salud del piloto. Que dos variables de nivel estructural — tipo de equipo de debut y longevidad en el deporte — expliquen el 41.6% de la varianza ya es informativo. Los coeficientes son significativos, la dirección es la esperada, y el modelo cumple su propósito descriptivo.')

    separador(doc)

    pregunta(doc, 4,
        '"¿Por qué usaron equal_var=False? ¿Verificaron el supuesto de varianzas?"',
        'El test de Welch con equal_var=False es hoy el estándar recomendado por defecto en la mayoría de los textos de estadística aplicada, precisamente porque es más robusto cuando las varianzas son desiguales y no pierde mucha potencia cuando sí lo son. En los boxplots se observa que las dispersiones de ambos grupos son distintas (~1.50 vs ~0.85), lo que sugiere que Welch es la elección correcta. Podríamos verificarlo formalmente con un test de Levene o Bartlett, algo que dejamos como mejora para el avance final.')

    separador(doc)

    pregunta(doc, 5,
        '"¿Verificaron los supuestos del t-test? ¿Normalidad, independencia?"',
        'El supuesto de normalidad es menos crítico en este caso porque ambos grupos tienen tamaños de muestra suficientemente grandes (n=123 y n=254) para invocar el teorema central del límite, que garantiza que la distribución de las medias muestrales se aproxima a la normal. El supuesto de independencia entre observaciones se cumple razonablemente: cada piloto es una observación independiente, no hay medidas repetidas sobre el mismo individuo.')

    separador(doc)

    pregunta(doc, 6,
        '"¿Por qué el coeficiente de es_top es 0.52 si la diferencia de medias es 0.99?"',
        'Excelente pregunta. La diferencia de medias del EDA (0.99 pts/carrera) es la diferencia bruta entre grupos, sin controlar nada. El coeficiente de la regresión (0.52) es el efecto de debutar en Top manteniendo constante el número de carreras. Al incluir n_carreras en el modelo, parte de la diferencia que atribuíamos al tipo de equipo resulta explicada por el hecho de que los pilotos de equipos Top tienden a tener carreras más largas. El coeficiente 0.52 es la estimación más rigurosa del efecto del debut.')

    # ── GUARDAR ───────────────────────────────────────────────────
    output = r'C:\Users\AndyV\analisispt2\Guion_Presentacion.docx'
    doc.save(output)
    print(f'Guion generado: {output}')

if __name__ == '__main__':
    build()
