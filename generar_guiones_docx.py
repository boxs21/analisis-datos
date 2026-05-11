# pip install python-docx
# python generar_guiones_docx.py
# Genera: Guion_Andy.docx  y  Guion_Javier.docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED  = RGBColor(0xE1, 0x06, 0x00)
BLUE = RGBColor(0x1A, 0x6B, 0xD0)
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GRAY = RGBColor(0x55, 0x55, 0x55)
WHIT = RGBColor(0xFF, 0xFF, 0xFF)

# ── helpers ────────────────────────────────────────────────────────────────

def sp(para, before=0, after=80, line=276):
    pPr = para._p.get_or_add_pPr()
    s = OxmlElement('w:spacing')
    s.set(qn('w:before'), str(before)); s.set(qn('w:after'), str(after))
    s.set(qn('w:line'), str(line));     s.set(qn('w:lineRule'), 'auto')
    pPr.append(s)

def body(doc, text, size=11.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sp(p, before=0, after=90, line=330)
    parts = text.split('**')
    for i, part in enumerate(parts):
        if not part:
            continue
        r = p.add_run(part)
        r.font.size = Pt(size)
        r.bold = (i % 2 == 1)
    return p

def nota(doc, text):
    p = doc.add_paragraph()
    sp(p, before=30, after=60, line=276)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement('w:ind'); ind.set(qn('w:left'), '440')
    pPr.append(ind)
    r = p.add_run(f'[Nota: {text}]')
    r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY

def sep(doc):
    p = doc.add_paragraph()
    sp(p, before=60, after=60, line=276)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('─ ─ ─')
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC); r.font.size = Pt(10)

def slide_header(doc, num, titulo, tiempo, color):
    p = doc.add_paragraph()
    sp(p, before=280, after=20, line=276)
    r1 = p.add_run(f'Slide {num}  '); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = color
    r2 = p.add_run(titulo);           r2.bold = True; r2.font.size = Pt(12); r2.font.color.rgb = DARK
    r3 = p.add_run(f'   ({tiempo})'); r3.font.size = Pt(10); r3.font.color.rgb = GRAY

def pregunta(doc, num, qtext, atext):
    p = doc.add_paragraph()
    sp(p, before=220, after=50, line=276)
    r = p.add_run(f'P{num}: {qtext}'); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK
    body(doc, atext)

def new_doc(nombre, color_acento):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.5)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    return doc

def portada(doc, nombre, slides_rango, color):
    doc.add_paragraph(); doc.add_paragraph()
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(titulo, before=0, after=80, line=360)
    r = titulo.add_run('GUION DE PRESENTACIÓN')
    r.bold = True; r.font.size = Pt(22); r.font.color.rgb = color

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(sub, before=40, after=200, line=276)
    r2 = sub.add_run(nombre.upper())
    r2.bold = True; r2.font.size = Pt(16); r2.font.color.rgb = DARK

    for label, value in [
        ('Proyecto:', '¿Debutar en un equipo grande te da más puntos en la F1?'),
        ('Curso:', 'Análisis de Datos e Inferencia Estadística — UDD'),
        ('Slides asignados:', slides_rango),
        ('Duración total:', '10 minutos aprox.'),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp(p, before=30, after=30, line=276)
        rl = p.add_run(label + '  '); rl.bold = True; rl.font.size = Pt(11)
        rv = p.add_run(value); rv.font.size = Pt(11)

    doc.add_paragraph()
    aviso = doc.add_paragraph()
    aviso.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(aviso, before=80, after=40, line=276)
    ra = aviso.add_run('Los textos entre corchetes son acotaciones escénicas — NO se leen en voz alta.')
    ra.italic = True; ra.font.size = Pt(10); ra.font.color.rgb = GRAY
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# GUION ANDY  — Slides 1–7  (~4 min 35 seg)
# ══════════════════════════════════════════════════════════════════════════════

def build_andy():
    doc = new_doc('Andy Villarroel', RED)
    portada(doc, 'Andy Villarroel', 'Slides 1 a 7', RED)

    # ── SLIDE 1 — Portada (0:30) ───────────────────────────────────
    slide_header(doc, 1, 'Portada', '0:30', RED)
    body(doc, 'Buenas tardes. Somos Andy Villarroel y Javier Alcaino. En los próximos diez minutos les presentamos nuestro análisis sobre Fórmula 1: **¿debutar en un equipo grande te da más puntos?** Les contaremos qué pregunta nos hicimos, cómo construimos el análisis y qué encontramos.')
    nota(doc, 'Andy señala el título. Javier asiente al lado. Pausa antes de pasar.')
    sep(doc)

    # ── SLIDE 2 — Agenda (0:20) ────────────────────────────────────
    slide_header(doc, 2, 'Agenda', '0:20', RED)
    body(doc, 'La presentación tiene cuatro bloques: pregunta e introducción, dataset y limpieza, análisis exploratorio, y finalmente Javier con la inferencia estadística y las conclusiones.')
    nota(doc, 'Recorrer los ítems visualmente. No leer cada uno palabra por palabra.')
    sep(doc)

    # ── SLIDE 3 — Pregunta (0:50) ──────────────────────────────────
    slide_header(doc, 3, 'Pregunta de Investigación', '0:50', RED)
    body(doc, 'Todo análisis parte de una pregunta concreta. La nuestra es: **¿el promedio de puntos por carrera —PPC— de pilotos que debutaron en equipos Top es significativamente mayor que el de pilotos de equipos Chico?**')
    body(doc, 'Definimos **equipo Top** como cualquier constructor que haya ganado el Campeonato Mundial de Constructores al menos una vez: Ferrari, McLaren, Mercedes, Red Bull, Williams, Lotus, entre otros. El **PPC** es simplemente los puntos totales dividido el número de carreras disputadas.')
    body(doc, 'La hipótesis nula es que no hay diferencia. La alternativa es que los pilotos de equipos Top tienen mayor PPC. En un momento Javier explica el test formal.')
    nota(doc, 'Señalar la fórmula PPC = puntos_totales / n_carreras en la slide. Tomarse este slide con calma, el profe prestará atención aquí.')
    sep(doc)

    # ── SLIDE 4 — Dataset (0:45) ───────────────────────────────────
    slide_header(doc, 4, 'Dataset', '0:45', RED)
    body(doc, 'Usamos el dataset de **Ergast Motor Racing**, API pública con datos de F1 desde 1950 hasta 2023. Es el dataset de referencia en análisis de este deporte.')
    body(doc, 'Trabajamos con cinco tablas: resultados, carreras, pilotos, constructores y clasificaciones de constructores. Después de la limpieza, el dataset final tiene **377 pilotos**: **123 del grupo Top** y **254 del grupo Chico**.')
    nota(doc, 'Señalar la tabla de fuentes y los dos bloques de distribución de pilotos en la slide.')
    sep(doc)

    # ── SLIDE 5 — Limpieza (0:35) ──────────────────────────────────
    slide_header(doc, 5, 'Limpieza de Datos', '0:35', RED)
    body(doc, 'Los datos crudos requirieron varias decisiones. La más importante: **filtramos pilotos con menos de 10 carreras**. Un piloto con dos o tres carreras podría tener un PPC de 10 puntos por haber ganado en su debut — eso no refleja rendimiento real, es ruido estadístico.')
    body(doc, 'También identificamos el equipo de debut de cada piloto, definimos Top mediante las clasificaciones de constructores, y los valores nulos en puntos se imputaron como cero — retiros sin puntaje.')
    nota(doc, 'Mantener ritmo, no detenerse demasiado aquí salvo que el profe pregunte.')
    sep(doc)

    # ── SLIDE 6 — EDA Descriptiva (0:45) ───────────────────────────
    slide_header(doc, 6, 'EDA — Estadística Descriptiva', '0:45', RED)
    body(doc, 'Los números ya cuentan una historia clara. El grupo Top tiene un **PPC promedio de 1.62 puntos por carrera**. El grupo Chico tiene **0.63 puntos por carrera**. Una diferencia de casi **1 punto por cada carrera**.')
    body(doc, 'En el campeonato moderno, el margen entre el primero y el segundo del WDC suele ser de 10 a 30 puntos en toda la temporada. Acumular 1 punto extra por carrera durante 100 o 200 Grandes Premios es una ventaja estructural enorme, no marginal.')
    nota(doc, 'Señalar los valores de la tabla y las barras comparativas en la slide.')
    sep(doc)

    # ── SLIDE 7 — Histograma (0:30) ────────────────────────────────
    slide_header(doc, 7, 'Distribución del PPC', '0:30', RED)
    body(doc, 'El histograma confirma lo que vimos en los números: ambas distribuciones son asimétricas hacia la derecha, pero la distribución del grupo Top está claramente corrida. La mayoría del grupo Chico se acumula entre 0 y 1 pto/carrera, mientras el grupo Top tiene más masa en rangos intermedios y altos.')
    body(doc, '**[Transición a Javier]:** Con ese panorama descriptivo, le paso la palabra a Javier para las pruebas estadísticas formales.')
    nota(doc, 'Señalar la diferencia de centros entre ambas distribuciones. Mirar a Javier al hacer la transición.')
    sep(doc)

    doc.add_page_break()

    # ── PREGUNTAS ANTICIPADAS ──────────────────────────────────────
    p = doc.add_paragraph()
    sp(p, before=0, after=100, line=276)
    r = p.add_run('PREGUNTAS ANTICIPADAS')
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK

    aviso = doc.add_paragraph()
    sp(aviso, before=0, after=120, line=276)
    ra = aviso.add_run('Esta sección es solo para preparación — no se lee durante la presentación.')
    ra.italic = True; ra.font.size = Pt(10); ra.font.color.rgb = GRAY

    sep(doc)

    pregunta(doc, 1,
        '"¿Por qué no estandarizaron los puntos según la época?"',
        'Es la principal limitación que reconocemos. No estandarizamos porque hacerlo requería definir un denominador de normalización con supuestos propios — el máximo posible por temporada, la posición relativa, etc. Optamos por reportar la limitación con honestidad. Para el avance final, usaríamos el porcentaje de puntos obtenidos sobre el máximo posible en cada temporada.')

    sep(doc)

    pregunta(doc, 2,
        '"¿Por qué filtrar solo pilotos con 10 o más carreras?"',
        'El umbral de 10 carreras es conservador pero funciona: elimina pilotos con PPC no representativo sin reducir demasiado el tamaño de muestra. Pilotos con menos de 10 carreras representarían outliers estructurales — un piloto que ganó su única carrera tendría PPC de 25, que es un ruido enorme. Podríamos haber usado 15 o 20, pero eso dependería de un análisis de sensibilidad que dejamos para el avance final.')

    sep(doc)

    pregunta(doc, 3,
        '"¿Cómo saben que el equipo causó el éxito y no el talento del piloto?"',
        'No lo sabemos, y somos explícitos en eso. Nuestro diseño es observacional — no podemos afirmar causalidad. Es una correlación estadísticamente significativa. Los equipos Top contratan a los pilotos percibidos como mejores, así que parte del mayor PPC puede reflejar talento previo, no solo el efecto del equipo. Para separar eso necesitaríamos un diseño cuasi-experimental, que está fuera del alcance de este trabajo.')

    doc.save(r'C:\Users\AndyV\analisispt2\Guion_Andy.docx')
    print('Guion_Andy.docx generado.')


# ══════════════════════════════════════════════════════════════════════════════
# GUION JAVIER  — Slides 8–14  (~5 min 25 seg)
# ══════════════════════════════════════════════════════════════════════════════

def build_javier():
    doc = new_doc('Javier Alcaino', BLUE)
    portada(doc, 'Javier Alcaino', 'Slides 8 a 14', BLUE)

    # ── SLIDE 8 — Boxplot y Scatter (0:50) ────────────────────────
    slide_header(doc, 8, 'Boxplot y Scatter', '0:50', BLUE)
    body(doc, 'Gracias, Andy. Antes de entrar al test, dos visualizaciones más.')
    body(doc, 'El **boxplot** confirma lo que Andy describió: la mediana del grupo Top está notablemente más alta que la del grupo Chico, con mayor IQR y outliers superiores que corresponden a las leyendas del deporte — Hamilton, Schumacher, Senna.')
    body(doc, 'El **scatter** —PPC versus número de carreras— muestra una tendencia positiva leve: los pilotos con más carreras tienden a tener PPC más altos. Eso me indica que el número de carreras es una covariable relevante que debo controlar en la regresión.')
    nota(doc, 'Señalar boxplot primero, luego el scatter. Transición fluida entre ambos.')
    sep(doc)

    # ── SLIDE 9 — Test de Hipótesis (1:05) ────────────────────────
    slide_header(doc, 9, 'Test de Hipótesis', '1:05', BLUE)
    body(doc, 'Nuestras hipótesis: **H₀** — el PPC promedio es igual en ambos grupos. **H₁** — el PPC del grupo Top es mayor.')
    body(doc, 'Usamos el **test T de Welch** —T con varianzas desiguales— porque los grupos tienen tamaños distintos (123 vs 254) y dispersiones distintas (~1.50 vs ~0.85). El test de Welch corrige los grados de libertad para este caso y es más robusto que el T estándar. En la implementación usamos `equal_var=False`.')
    body(doc, 'El resultado fue contundente: **T = 4.61**, **p-valor = 7.86 × 10⁻⁶** — varios órdenes de magnitud por debajo de 0.05. **Rechazamos H₀**. Hay evidencia estadística sólida de que los pilotos que debutaron en equipos Top tienen PPC significativamente mayor.')
    body(doc, 'Ahora bien, rechazar H₀ no nos dice cuánto mayor es la diferencia ni si hay otras variables involucradas. Para eso, el modelo de regresión.')
    nota(doc, 'Señalar el p-valor en la slide. Si el profe pregunta bilateral vs unilateral: el código usa ttest_ind que es bilateral por defecto — para test unilateral el p sería la mitad, aún más significativo.')
    sep(doc)

    # ── SLIDE 10 — Regresión OLS (0:50) ───────────────────────────
    slide_header(doc, 10, 'Modelo de Regresión OLS', '0:50', BLUE)
    body(doc, 'Construí un modelo OLS. Variable dependiente: el PPC. Variables independientes: **es_top** —dummy 1 si debutó en Top, 0 si no— y **n_carreras** como control de experiencia.')
    body(doc, 'Incluí n_carreras para aislar el efecto del debut: quiero que la diferencia no se explique simplemente porque los pilotos de equipos Top tienden a tener carreras más largas.')
    body(doc, 'El modelo tiene un **R² de 41.6%**. Ambos coeficientes son significativos con p < 0.001. La constante también es significativa (p ≈ 0.04). Les detallo los coeficientes en el siguiente slide.')
    nota(doc, 'Señalar la tabla de coeficientes y el R² en la slide.')
    sep(doc)

    # ── SLIDE 11 — Interpretación de Coeficientes (0:50) ──────────
    slide_header(doc, 11, 'Interpretación de Coeficientes', '0:50', BLUE)
    body(doc, '**Constante β₀ ≈ −0.21**: PPC teórico para un piloto de equipo Chico con cero carreras. Al ser negativo y fuera del rango real de datos, no tiene interpretación práctica — solo ancla matemáticamente el modelo.')
    body(doc, '**Coeficiente es_top ≈ +0.52 — el más importante**: controlando por el número de carreras, debutar en Top agrega **0.52 pts/carrera al PPC**. Este 0.52 es distinto a la diferencia bruta de 0.99 que Andy mostró en el EDA: la regresión descuenta el efecto de la experiencia, y el efecto neto del debut es 0.52.')
    body(doc, '**Coeficiente n_carreras ≈ +0.015**: cada carrera adicional suma 0.015 pts al PPC. Pequeño por carrera, pero 200 carreras equivalen a ~3 pts extra de PPC — los pilotos más longevos tienden a ser más consistentes.')
    nota(doc, 'Señalar los coeficientes en la tabla. Marcar bien la diferencia entre 0.52 (OLS controlado) y 0.99 (diferencia bruta del EDA).')
    sep(doc)

    # ── SLIDE 12 — Discusión y Limitaciones (0:55) ────────────────
    slide_header(doc, 12, 'Discusión y Limitaciones', '0:55', BLUE)
    body(doc, 'Los resultados son claros estadísticamente: diferencia real y robusta tanto en el t-test como en la regresión. Pero la interpretación causal es delicada.')
    body(doc, '**Limitación principal**: el sistema de puntos cambió radicalmente. En los años 50 el ganador se llevaba 9 puntos; hoy se llevan 25. Un piloto de los 60 y uno de 2020 no son comparables en PPC sin estandarización. Eso es trabajo futuro.')
    body(doc, '**Segunda limitación**: sesgo de selección. Ferrari y Mercedes contratan a los pilotos percibidos como los mejores. Parte del mayor PPC refleja talento previo, no solo el equipo. No podemos separarlo con este diseño observacional.')
    nota(doc, 'Hablar con seguridad sobre las limitaciones — mencionarlas muestra pensamiento crítico, el profe valora eso.')
    sep(doc)

    # ── SLIDE 13 — Próximos Pasos (0:25) ──────────────────────────
    slide_header(doc, 13, 'Próximos Pasos', '0:25', BLUE)
    body(doc, 'Tres extensiones naturales: **primero**, estandarizar el sistema de puntos por era. **Segundo**, analizar la trayectoria completa del piloto, no solo el debut. **Tercero**, verificar los supuestos del OLS con QQ-plot y test de Breusch-Pagan, y explorar bootstrap para estimación no paramétrica.')
    nota(doc, 'Ir rápido en este slide.')
    sep(doc)

    # ── SLIDE 14 — Cierre (0:30) ───────────────────────────────────
    slide_header(doc, 14, 'Conclusiones y Cierre', '0:30', BLUE)
    body(doc, 'Para cerrar: encontramos evidencia estadística sólida de que los pilotos que debutan en equipos campeones del mundo acumulan significativamente más puntos por carrera. La diferencia bruta es de 0.99 pts/carrera; el efecto controlado por la regresión es **0.52 pts/carrera**. Ambos son estadísticamente robustos.')
    body(doc, '"Si Hamilton hubiera debutado en un Minardi, ¿dónde estaría hoy? La respuesta estadística sería: con un PPC bastante más bajo. Pero eso es especulación — y nosotros somos gente de datos."')
    body(doc, '**Muchas gracias. Quedamos disponibles para preguntas.**')
    nota(doc, 'Ambos al frente. Javier termina. Dejar silencio breve antes de abrir preguntas.')
    sep(doc)

    doc.add_page_break()

    # ── PREGUNTAS ANTICIPADAS ──────────────────────────────────────
    p = doc.add_paragraph()
    sp(p, before=0, after=100, line=276)
    r = p.add_run('PREGUNTAS ANTICIPADAS')
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK

    aviso = doc.add_paragraph()
    sp(aviso, before=0, after=120, line=276)
    ra = aviso.add_run('Esta sección es solo para preparación — no se lee durante la presentación.')
    ra.italic = True; ra.font.size = Pt(10); ra.font.color.rgb = GRAY

    sep(doc)

    pregunta(doc, 1,
        '"El R² es 41.6% — ¿no significa que el modelo no sirve?"',
        'El R² de 41.6% con solo dos variables es un resultado sólido para rendimiento deportivo humano. Estamos modelando el desempeño en un deporte donde el auto, el motor, la estrategia, el clima y la suerte introducen enorme variabilidad. Que dos variables estructurales expliquen el 41.6% de la varianza ya es informativo. Los coeficientes son significativos y la dirección es la esperada.')

    sep(doc)

    pregunta(doc, 2,
        '"¿Por qué el coeficiente es_top es 0.52 y no 0.99 como la diferencia de medias?"',
        'La diferencia de medias de 0.99 es la diferencia bruta sin controlar nada. El coeficiente 0.52 de la regresión es el efecto del debut manteniendo constante el número de carreras. Al incluir n_carreras en el modelo, parte de la diferencia entre grupos se explica porque los pilotos de equipos Top tienden a disputar más carreras. El coeficiente 0.52 es la estimación más rigurosa y conservadora del efecto del debut.')

    sep(doc)

    pregunta(doc, 3,
        '"¿Por qué usaron equal_var=False? ¿Verificaron el supuesto?"',
        'El test de Welch con equal_var=False es el estándar recomendado en estadística aplicada moderna porque es robusto cuando las varianzas son desiguales sin perder potencia cuando son iguales. Los boxplots muestran dispersiones distintas entre grupos (~1.50 vs ~0.85), lo que sugiere que Welch es la elección correcta. Podríamos verificarlo formalmente con un test de Levene — algo que incluiríamos en el avance final.')

    sep(doc)

    pregunta(doc, 4,
        '"¿Verificaron normalidad para el t-test?"',
        'Con muestras de n=123 y n=254, el Teorema Central del Límite garantiza que las distribuciones muestrales de las medias se aproximan a la normal independientemente de la distribución original. El t-test de Welch es robusto a desviaciones de normalidad en muestras grandes. El supuesto de independencia se cumple razonablemente: cada piloto es una observación independiente.')

    sep(doc)

    pregunta(doc, 5,
        '"¿La constante negativa no invalida el modelo?"',
        'No. La constante negativa (-0.21) simplemente indica que el valor teórico de PPC para un piloto con cero carreras y debut en equipo Chico es negativo — algo imposible en la práctica, pero eso es porque el modelo extrapola fuera del rango real de los datos. La constante solo ancla el modelo; lo que importa son los coeficientes de las variables de interés, que son positivos, significativos y con sentido sustantivo.')

    doc.save(r'C:\Users\AndyV\analisispt2\Guion_Javier.docx')
    print('Guion_Javier.docx generado.')


# ══════════════════════════════════════════════════════════════════════════════

if __name__ == '__main__':
    build_andy()
    build_javier()
    print('\nTiempos estimados:')
    print('  Andy   (slides 1-7):  ~4 min 35 seg')
    print('  Javier (slides 8-14): ~5 min 25 seg')
    print('  TOTAL:                ~10 min 00 seg')
